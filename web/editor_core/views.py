from django.shortcuts import render
import os,re
import subprocess, shlex, json
from django.http.response import HttpResponse
from . import get_help
from django.http.response import FileResponse
from . import beauty_dir,file_unzip
from JavaParser_Web import settings

# Create your views here.

def changepdf():
    import comtypes.client
    import time
    import os
    import pythoncom,  win32com.client
    # Variables and Inputs=================================================
    root = os.getcwd().replace('\\','/')

    File = root + '/upload/demo.docx'  # Or .doc, rtf files
    outFile = root + '/upload/demo.pdf'

    print(outFile)


    # Functions ============================================================
    def convert_word_to_pdf(inputFile, outputFile):
        ''' the following lines that are commented out are items that others shared with me they used when
        running loops to stop some exceptions and errors, but I have not had to use them yet (knock on wood) '''
        print('start')
        # word.visible = True
        word = win32com.client.Dispatch('Word.Application')

        doc = word.Documents.Open(inputFile, ReadOnly = 1)
        doc.SaveAs(outputFile, FileFormat=17)
        doc.close()
        # word.visible = False
        word.Quit()

    # Main Body=================================================================
    convert_word_to_pdf(File, outFile)

def py_doc(dict):
    try:
        os.remove(os.path.join('upload', 'demo.docx'))
        os.remove(os.path.join('upload', 'demo.pdf'))
    except Exception as e:
        print(e)
    from docx import Document
    from docx.shared import Inches
    import json
    # print('生成docx调试',dict)
    from docx.oxml.ns import qn



    try:
        beauty_code = dict['fixedCode']
        problems = dict['problems']
        import random
        score = 100 - len(problems)*5 + random.random()
    except:
        score = 0

    document = Document()

    document.add_heading('Static Analysis Report', 0)

    p = document.add_paragraph('the analyze for your ')
    p.add_run('code').bold = True
    p.add_run(' and some ')
    p.add_run('details.').italic = True
    document.add_heading('Your Code is about:' + str(round(score, 1)) + "points", level=1)
    document.add_heading('Problems', level=1)

    document.styles['Normal'].font.name = u'Consolas'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'Consolas')

    if(dict['success'] == False):
        for err in dict['error']:
            err = err.split('\r\n')[0]
            document.add_paragraph(err, style='Intense Quote')

    else:
        for p in problems:
            try:
                id = p.split(':')[1]
            except:
                id = p
            try:
                hint = get_help.get_help('#'+id)
                from pyquery import PyQuery as pq
                doc = pq(hint)
                hint = doc.text()
                p = p + '\n' + hint
            except Exception as e:
                print(e)
            document.add_paragraph(p, style='Intense Quote')
        func_def = dict['originCode']
        print('源代码正常',func_def)

        document.add_heading('Your Origin Code', level=1)
        document.add_paragraph(func_def)

        document.add_heading('Beauty_Code', level=1)
        document.add_paragraph(beauty_code.replace('\r', '\n').replace('\r\n', '\n').replace('\n\n', '\n'))

    junk = "Attribution — You must give appropriate credit, provide a link to the license, and indicate if changes were made. You may do so in any reasonable manner, but not in any way that suggests the licensor endorses you or your use."

    junk2 = "NonCommercial — You may not use the material for commercial purposes."

    junk3 = "NoDerivatives — If you remix, transform, or build upon the material, you may not distribute the modified material."

    junk4 = "No additional restrictions — You may not apply legal terms or technological measures that legally restrict others from doing anything the license permits."
    document.add_heading('Under the following terms:', level=1)
    document.add_paragraph(junk, style='Intense Quote')
    document.add_paragraph(junk2, style='Intense Quote')
    document.add_paragraph(junk3, style='Intense Quote')
    document.add_paragraph(junk4, style='Intense Quote')


    document.save('upload/demo.docx')
    try:
        os.system('soffice --headless --convert-to pdf upload/demo.docx && copy demo.pdf upload')
        # changepdf()
    except Exception as e:
        print(e)
        print('pdf转换失败')

def index(request):
    # core_value = request.POST['core']
    # core_id = request.POST['id']
    core_id = 'editor1'
    with open('editor_core/editor.txt','r') as file:
        text = file.read()
    # with open('editor_core/ans.txt','r') as file:
    #     core_value = file.read()
    text = text.replace('editor_name', 'editor_1')
    text = text.replace('core_id',core_id)
    text = text.replace('core_value', 'input_your_code')
    text = text.replace('read_choose', 'false')

    context = {}
    context['core_editor'] = text

    core_id = 'editor2'
    with open('editor_core/editor.txt','r') as file:
        text2 = file.read()
    text2 = text2.replace('core_id',core_id)
    text2 = text2.replace('core_value', 'beauty_code')
    text2 = text2.replace('read_choose', 'true')
    text2 = text2.replace('editor_name', 'editor_2')
    context['core_editor_beauty'] = text2

    return render(request, 'index.html', context)

def test_index(request):
    return render(request, 'test_index.html')

def file_download_report(request):

    file = open(os.path.join('upload', 'demo.pdf'), 'rb')

    response = FileResponse(file)
    response['Content-Type'] = 'application/octet-stream'
    response['Content-Disposition'] = 'attachment;filename="report.pdf"'

    return response


def beautiful_core(request):

    choose1 = request.POST.get('choose1')
    choose2 = request.POST.get('choose2')
    choose3 = request.POST.get('choose3')

    choose4 = request.POST.get('choose4')

    print(choose1,choose2,choose3,choose4)

    core_value = request.POST.get('core')
    # print(core_value)

    with open('editor_core/junkjava.txt', 'w+',newline='') as file:
        file.write(core_value)

    add = ""

    if (str(choose4) == 'true'):
        list = {}
        import ast
        with open('editor_core/junkjava.txt', 'r',) as file:
            text = file.read()
            list['originCode'] = text;
        try:
            r_node = ast.parse(text)
            # print(astunparse.dump(r_node))  # 打印具体的抽象语法树,如果整个语法树有错误，则识别失败
        except Exception as e:
            reason = e.args[1][0] + e.args[1][3]
            line1 = str(e.args[1][1])
            line2 = str(e.args[1][2])
            text = '(line 9,col 5) '.replace('9', line1).replace('5', line2) + reason
            list = {}
            t = []
            t.append(text)
            list['success'] = False
            list['error'] = t
            return HttpResponse(json.dumps(list), content_type='application/json,charset=utf-8')

        problem = os.popen('pycodestyle --first editor_core/junkjava.txt').read()

        problem = str(problem).replace('editor_core/junkjava.txt:','')

        problem = str(problem).replace(': ', '@')
        problem = str(problem).replace(':', ',')
        problem = str(problem).replace('@', ': ')

        p = []

        for str_spt in problem.split('\n'):
            try:
                str_spt = str(str_spt)
                print(str_spt)
                tmp = str_spt.split(',')
                print(tmp)
                a = '(line ' + str_spt.split(',')[0] + ',col '
                print(a)
                b = str_spt.split(',')[1].split(':')[0] + ')'
                c =  a + b+ '-'+ a +b
                text = str_spt.split(',')[1].split(':')[1]
                print(text)
                anstext = ""
                cnt = 0
                for itm in text.split(' '):
                    if(cnt == 0):
                        cnt += 1
                    else:
                        anstext += itm
                        anstext += ' '
                p.append(c+' '+ anstext)
            except Exception as e:
                print(e)


        print(problem)

        os.system('autopep8 --in-place  editor_core/junkjava.txt')
        with open('editor_core/junkjava.txt', 'r',) as file:
            text = file.read()




        list['success'] = True
        list['fixedCode'] = text
        list['problems'] = p

        # return_json = json.dumps(list).replace('"','\\'+'"')
        # return_json = '"' + return_json + '"'

        return_json = json.dumps(list)
        print(return_json)

        try:
            py_doc(list)
        except Exception as e:
            print(e)
        return HttpResponse(return_json, content_type='application/json,charset=utf-8')

    else:

        if(str(choose1) == 'true'):
            add += '  --expandSingleIf'
        if(str(choose2) == 'true'):
            add += '  --transferSwitchToIf'
        if(str(choose3) == 'true'):
            add += '  --transferForToWhile'

        command = 'java -cp JavaIntelligence.jar org.hacksource.cli.IntelligenceCLI' + add
        print(command)

        args = shlex.split(command)
        p = subprocess.Popen(args, stdin=open('editor_core/junkjava.txt', 'r'), stdout=subprocess.PIPE)
        print('.....-1')
        output, unused_err  = p.communicate()
        print('.....0')
        p2 = str(output, encoding='utf-8')
        print('p2=',p2)
        return_dict  = (json.loads(p2))
        return_json = (json.dumps(p2))
        print(return_json)
        try:
            py_doc(return_dict)
        except Exception as e:
            print(e)

        return HttpResponse(return_json, content_type='application/json,charset=utf-8')






def get_help_hint(request):
    help_id = request.POST.get('help_id')
    hint = get_help.get_help(help_id)
    return HttpResponse(hint)

def acceptfiles(request):
    obj = request.FILES.get('fafafa')

    f = open(os.path.join('upload', 'beauty_dir.zip'), 'wb')
    for line in obj.chunks():
        f.write(line)
    f.close()

    try:
        os.remove(os.path.join('upload', 'beauty_finished.zip'))
    except Exception as e:
        print(e)
    file_unzip.handlezip()
    file = open(os.path.join('upload', 'beauty_finished.zip'), 'rb')

    response = FileResponse(file)
    response['Content-Type'] = 'application/octet-stream'
    response['Content-Disposition'] = 'attachment;filename="beauty_finished.zip"'

    return response

    # return HttpResponse('上传成功')


def file_download(request):
    try:
        os.remove(os.path.join('upload', 'beauty_finished.zip'))
    except Exception as e:
        print(e)

    file_unzip.handlezip()
    file = open(os.path.join('upload', 'beauty_finished.zip'), 'rb')

    response = FileResponse(file)
    response['Content-Type'] = 'application/octet-stream'
    response['Content-Disposition'] = 'attachment;filename="beauty_finished.zip"'



    return response

